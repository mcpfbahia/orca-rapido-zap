import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import RGBColor, Pt, Cm
import os
import datetime
import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

st.markdown("""
<div style='text-align: center; padding: 12px 0; border-bottom: 2px solid #ccc;'>
    <h1 style='margin-bottom: 5px; font-size: 36px; color: black; background-color: #f9c922; display: inline-block; padding: 8px 16px; border-radius: 8px;'>
        Proposta Personalizada MCPF BAHIA
    </h1>
    <p style='margin-top: 8px; font-size: 18px; color: #333;'>Madeira Tratada com Garantia</p>
</div>
""", unsafe_allow_html=True)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
kits_file = os.path.join(BASE_DIR, 'kits.xlsx')
modelo_default = os.path.join(BASE_DIR, 'modelo-so-kit.docx')
output_dir = os.path.join(BASE_DIR, 'propostas_geradas')

def formatar_moeda(valor):
    if valor is None or valor == '' or (isinstance(valor, float) and pd.isna(valor)):
        return "Cálculo para o modelo não gerado"
    try:
        return f"R$ {float(valor):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except Exception:
        return "Cálculo para o modelo não gerado"

def slugify(value):
    value = str(value).strip().replace(" ", "_")
    value = re.sub(r'[^\w\-_\.]', '', value)
    return value

def aplicar_negrito(paragrafo, substituicoes):
    texto_original = paragrafo.text
    for chave, valor in substituicoes.items():
        texto_original = texto_original.replace(chave, f"§§§{valor}§§§")
    if texto_original == paragrafo.text:
        return
    partes = re.split(r'(§§§.*?§§§|R\$ [\d\.,]+)', texto_original)
    paragrafo.clear()
    for parte in partes:
        run = paragrafo.add_run(parte.replace("§§§", ""))
        if parte.startswith("§§§") and parte.endswith("§§§"):
            run.bold = True
        elif re.match(r'R\$ [\d\.,]+', parte):
            run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

def padronizar_fonte(modelo, nome_fonte="Segoe UI", tamanho=12):
    for p in modelo.paragraphs:
        for run in p.runs:
            run.font.name = nome_fonte
            run.font.size = Pt(tamanho)
    for tabela in modelo.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for par in celula.paragraphs:
                    for run in par.runs:
                        run.font.name = nome_fonte
                        run.font.size = Pt(tamanho)

def ajustar_tabela_fonte(tabela, nome_fonte="Segoe UI", tamanho=11):
    for row in tabela.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(tamanho)
                    run.font.name = nome_fonte

def extrair_area(descricao, campo_area):
    if not pd.isnull(campo_area):
        try:
            valor = float(str(campo_area).replace(',', '.'))
            return valor
        except:
            pass
    if isinstance(descricao, str):
        padrao = re.search(r'(\d+[\.,]?\d*)\s*(m2|m²)', descricao, re.IGNORECASE)
        if padrao:
            valor = float(padrao.group(1).replace(',', '.'))
            return valor
    return 0.0

# --- NOVA REGRA CHAVE NA MÃO ---
def calcular_chave_na_mao(descricao, area):
    desc = str(descricao).lower()
    adicionais = [
        "stain", "telha", "forro", "assoalho", "parede dupla",
        "externo", "impregnante"
    ]
    if re.search(r"camping\s*1", desc):
        return area * 2200
    elif re.search(r"camping\s*2", desc):
        return area * 2400
    elif re.search(r"camping\s*3", desc):
        return area * 2400
    elif "a-frame" in desc or "aframe" in desc:
        if area <= 60:
            return area * 1700
        else:
            return area * 1650
    elif ("kit" in desc and not any(x in desc for x in ["camping", "a-frame", "aframe"])
          and not any(adicional in desc for adicional in adicionais)):
        if area <= 42:
            return area * 2000
        else:
            return area * 1900
    elif ("pop" in desc or "pousada pop" in desc or "tiny house" in desc) and not any(adicional in desc for adicional in adicionais):
        if area <= 42:
            return area * 2000
        else:
            return area * 1900
    return None

def inserir_tabela_no_local(modelo, marcador, tabela_kits, total_geral):
    for i, par in enumerate(modelo.paragraphs):
        if marcador in par.text:
            p_element = par._element
            p_element.getparent().remove(p_element)
            tabela = modelo.add_table(rows=1, cols=6)  # Adiciona coluna para Chave na Mão!
            tabela.style = 'Light Grid'
            largura_colunas = [Cm(1.0), Cm(11), Cm(1.5), Cm(2.0), Cm(2.0), Cm(3.0)]
            for idx, width in enumerate(largura_colunas):
                tabela.columns[idx].width = width
            hdr_cells = tabela.rows[0].cells
            hdr_cells[0].text = "QUANT"
            hdr_cells[1].text = "KIT / DESCRIÇÃO"
            hdr_cells[2].text = "VALOR UNITÁRIO"
            hdr_cells[3].text = "VALOR TOTAL"
            hdr_cells[4].text = "VALOR C/ DESC."
            hdr_cells[5].text = "CHAVE NA MÃO"
            for row in tabela_kits:
                row_cells = tabela.add_row().cells
                for j, val in enumerate(row):
                    row_cells[j].text = str(val)
            total_row = tabela.add_row().cells
            total_row[0].text = ""
            total_row[1].text = "TOTAL"
            total_row[2].text = ""
            total_row[3].text = ""
            total_row[4].text = formatar_moeda(total_geral)
            total_row[5].text = ""
            ajustar_tabela_fonte(tabela, nome_fonte="Segoe UI", tamanho=11)
            tbl_element = tabela._element
            body = modelo._body._element
            if i < len(modelo.paragraphs) - 1:
                body.insert(i, tbl_element)
            else:
                body.append(tbl_element)
            break

def gerar_proposta_multikits(
    kits_file, modelo_file, lista_kits, nome_cliente, desconto_percentual, distancia_loja
):
    df_kits = pd.read_excel(kits_file)
    total_geral = 0
    total_peso = 0
    total_area = 0
    tabela_kits = []

    links_kits = []
    resumo_valores_kits = []
    total_valor_bruto = 0
    total_quantidade = 0
    total_desconto = 0

    # Prepara os dados para a tabela e resumos
    for item in lista_kits:
        modelo_selecionado = item["DESCRICAO"]
        quantidade = int(item["QUANTIDADE"])
        kit_row = df_kits[df_kits['DESCRICAO'] == modelo_selecionado].iloc[0]
        preco_normal = float(kit_row['A VISTA'])
        peso_unit = float(kit_row['PESO UND']) if not pd.isnull(kit_row['PESO UND']) else 0
        area_unit = extrair_area(str(kit_row['DESCRICAO']), kit_row['AREA'] if 'AREA' in kit_row.index else "")
        valor_kit = preco_normal * quantidade
        valor_avista = valor_kit * (1 - desconto_percentual / 100)
        desconto_kit = valor_kit - valor_avista

        # NOVO: calcula estimativa chave na mão para cada kit principal
        valor_chave_mao = calcular_chave_na_mao(modelo_selecionado, area_unit)
        valor_chave_mao_str = formatar_moeda(valor_chave_mao)

        tabela_kits.append([
            str(quantidade),
            modelo_selecionado,
            formatar_moeda(preco_normal),
            formatar_moeda(valor_kit),
            formatar_moeda(valor_avista),
            valor_chave_mao_str
        ])
        total_geral += valor_avista
        total_peso += peso_unit * quantidade
        total_area += area_unit * quantidade
        total_valor_bruto += valor_kit
        total_quantidade += quantidade
        total_desconto += desconto_kit

        link = kit_row['LINK_KIT']
        links_kits.append(f"{modelo_selecionado}: {link}")

        resumo_valores_kits.append(
            f"• {modelo_selecionado}: Valor Unit. {formatar_moeda(preco_normal)} / QTD: {quantidade} / Valor Total: {formatar_moeda(valor_kit)} / Desconto: {formatar_moeda(desconto_kit)} / Valor c/ Desc: {formatar_moeda(valor_avista)} / Chave na Mão: {valor_chave_mao_str}"
        )

    # Calcula frete e estimativas com base no total_peso e total_area
    frete_normal = (total_peso / 1000) * 1129
    distancia_ref = 200
    frete_adicional = max(0, (distancia_loja - distancia_ref)) * 5.50
    frete_total = frete_normal + frete_adicional
    valor_final_com_frete = total_geral + frete_total

    # Não calcula mais chave na mão total geral, apenas mostra por kit na tabela!

    cub_alvenaria = 2500
    cub_prefab = 1800
    custo_alvenaria = cub_alvenaria * total_area
    custo_chave_mao = cub_prefab * total_area
    economia_cub = custo_alvenaria - custo_chave_mao
    custo_mcpf = cub_prefab * total_area

    substituicoes = {
        '{{data_atual}}': datetime.datetime.today().strftime('%d/%m/%Y'),
        '{{NOME_CLIENTE}}': nome_cliente,
        '{{nome_cliente}}': nome_cliente,
        '{{area_total}}': f"{total_area:.2f} m²",
        '{{peso_total}}': f"{total_peso:.2f} kg",
        '{{valor_avista}}': formatar_moeda(total_geral),
        '{{frete_normal}}': formatar_moeda(frete_normal),
        '{{frete_adicional}}': formatar_moeda(frete_adicional),
        '{{frete_total}}': formatar_moeda(frete_total),
        '{{frete_total+valor_avista}}': formatar_moeda(valor_final_com_frete),
        '{{porcentagem_desconto}}': f"{desconto_percentual}%",
        '{{distancia_loja}}': f"{distancia_loja:.0f} km",
        '{{cub_alvenaria}}': formatar_moeda(cub_alvenaria),
        '{{cub_prefab}}': formatar_moeda(cub_prefab),
        '{{custo_alvenaria}}': formatar_moeda(custo_alvenaria),
        '{{custo_chave_mao}}': formatar_moeda(custo_chave_mao),
        '{{economia_cub}}': formatar_moeda(economia_cub),
        '{{custo_mcpf}}': formatar_moeda(custo_mcpf),
        # Chave na mão individual já vai na tabela
        '{{link_kit}}': "||LINK_PLACEHOLDER||",
        '{{resumo_valores_kits}}': "\n".join(resumo_valores_kits),
        '{{preço_normal}}': formatar_moeda(total_valor_bruto/total_quantidade) if total_quantidade else 'R$ 0,00',
        '{{quant}}': str(total_quantidade),
        '{{valor_total}}': formatar_moeda(total_valor_bruto),
        '{{desconto}}': formatar_moeda(total_desconto),
        '{{50%_valor_avista}}': formatar_moeda(total_geral/2),
        '{{area_casa}}': f"{total_area:.2f} m²",
    }

    modelo = Document(modelo_file)
    inserir_tabela_no_local(modelo, "{{modelo-selecionado}}", tabela_kits, total_geral)

    for p in modelo.paragraphs:
        aplicar_negrito(p, substituicoes)
    for tabela in modelo.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for par in celula.paragraphs:
                    aplicar_negrito(par, substituicoes)

    for i, p in enumerate(modelo.paragraphs):
        if '||LINK_PLACEHOLDER||' in p.text:
            p_element = p._element
            parent = p_element.getparent()
            idx = parent.index(p_element)
            parent.remove(p_element)
            for link in links_kits[::-1]:
                new_par = modelo.add_paragraph(link)
                parent.insert(idx, new_par._element)
            break

    padronizar_fonte(modelo, nome_fonte="Segoe UI", tamanho=12)

    os.makedirs(output_dir, exist_ok=True)
    nome_limpo = slugify(nome_cliente)
    output_path = os.path.join(output_dir, f"Proposta_{nome_limpo}.docx")
    modelo.save(output_path)
    return output_path

# ============= INTERFACE STREAMLIT =============

st.title("Proposta Comercial - Múltiplos Kits de Madeiramento")

nome_cliente = st.text_input("Nome do Cliente")
desconto_percentual = st.slider("Desconto (%)", min_value=0, max_value=12, value=5)
distancia_loja = st.number_input("Distância até o local da obra (km)", min_value=0.0, value=0.0, step=1.0)

if os.path.exists(kits_file):
    df_temp = pd.read_excel(kits_file)
    lista_modelos = df_temp['DESCRICAO'].tolist()

    if "kits_selecionados" not in st.session_state:
        st.session_state["kits_selecionados"] = []

    st.subheader("Adicionar Kit à Proposta")
    busca_kit = st.text_input("Buscar modelo de kit:")
    resultados_kit = [modelo for modelo in lista_modelos if busca_kit.lower() in modelo.lower()]
    kit_para_adicionar = st.selectbox("Selecione o modelo encontrado:", options=resultados_kit) if resultados_kit else None
    qtd_para_adicionar = st.number_input("Quantidade", min_value=1, step=1, value=1, key="qtd_para_adicionar")

    if st.button("Adicionar Kit"):
        if kit_para_adicionar:
            ja_adicionado = False
            for kit in st.session_state["kits_selecionados"]:
                if kit["DESCRICAO"] == kit_para_adicionar:
                    kit["QUANTIDADE"] += qtd_para_adicionar
                    ja_adicionado = True
                    break
            if not ja_adicionado:
                st.session_state["kits_selecionados"].append({
                    "DESCRICAO": kit_para_adicionar,
                    "QUANTIDADE": qtd_para_adicionar
                })
        else:
            st.warning("Selecione um kit para adicionar.")

    st.subheader("Kits Selecionados")
    if st.session_state["kits_selecionados"]:
        st.table(pd.DataFrame(st.session_state["kits_selecionados"]))
    else:
        st.info("Nenhum kit adicionado ainda.")

    if st.session_state["kits_selecionados"]:
        remove_index = st.number_input("Remover kit nº (índice da tabela, começando em 0)", min_value=0, max_value=len(st.session_state["kits_selecionados"])-1, step=1)
        if st.button("Remover Kit"):
            st.session_state["kits_selecionados"].pop(remove_index)
else:
    st.warning(f"⚠️ Arquivo {kits_file} não encontrado.")
    st.session_state["kits_selecionados"] = []

modelo_file = st.file_uploader("Modelo DOCX (opcional)", type=["docx"])

if st.button("📄 Gerar Proposta"):
    if not nome_cliente or not st.session_state["kits_selecionados"]:
        st.warning("Preencha o nome do cliente e adicione pelo menos um kit.")
    else:
        modelo = modelo_file if modelo_file else modelo_default
        try:
            caminho = gerar_proposta_multikits(
                kits_file,
                modelo,
                st.session_state["kits_selecionados"],
                nome_cliente,
                desconto_percentual,
                distancia_loja
            )
            if caminho:
                st.success("✅ Proposta gerada com sucesso!")
                with open(caminho, "rb") as file:
                    st.download_button("📥 Baixar Proposta", file, file_name=os.path.basename(caminho))
        except Exception as e:
            st.error(f"Erro ao gerar proposta: {type(e).__name__}: {e}")
